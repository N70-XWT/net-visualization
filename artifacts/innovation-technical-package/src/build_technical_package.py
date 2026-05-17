from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "artifacts" / "innovation-technical-package" / "output"
DOCX_PATH = OUT_DIR / "网络态势感知与可视化系统技术资料包.docx"


COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "navy": RGBColor(11, 37, 69),
    "muted": RGBColor(87, 101, 120),
    "body": RGBColor(30, 41, 59),
    "light_fill": "F2F4F7",
    "blue_fill": "E8EEF5",
    "soft_fill": "F4F6F9",
    "border": "D7DBE2",
    "risk": RGBColor(155, 28, 28),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DBE2", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def mark_first_row_as_header(table):
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, before=0, after=6, line=1.10, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", size=11, bold=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color or COLORS["body"])
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        before, after, size, color = 16, 8, 16, COLORS["blue"]
    elif level == 2:
        before, after, size, color = 12, 6, 13, COLORS["blue"]
    else:
        before, after, size, color = 8, 4, 12, COLORS["dark_blue"]
    style_paragraph(p, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    p.style = f"Heading {min(level, 3)}"
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if level:
        p.paragraph_format.left_indent = Inches(0.75)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=COLORS["body"])
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, before=0, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=COLORS["body"])
    return p


def add_table(doc, headers, rows, widths, header_fill=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_widths(table, widths)
    mark_first_row_as_header(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        if header_fill:
            set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.10, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=COLORS["navy"])

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.12, align=align)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=COLORS["body"])
    set_table_widths(table, widths)
    mark_first_row_as_header(table)
    add_para(doc, "", size=1, after=3)
    return table


def add_callout(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    mark_first_row_as_header(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=3, line=1.12)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=COLORS["navy"])
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0, line=1.12)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=COLORS["body"])
    add_para(doc, "", size=1, after=3)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 24, COLORS["navy"], 0, 10),
        ("Subtitle", 12, COLORS["muted"], 0, 12),
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    p = header.paragraphs[0]
    style_paragraph(p, after=0)
    r = p.add_run("中国国际大学生创新大赛技术资料包")
    set_run_font(r, size=9, color=COLORS["muted"])

    footer = section.footer
    p = footer.paragraphs[0]
    style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("网络态势感知与可视化系统")
    set_run_font(r, size=9, color=COLORS["muted"])


def add_cover(doc):
    add_para(doc, "技术资料包", size=12, bold=True, color=COLORS["blue"], after=4)
    title = doc.add_paragraph()
    style_paragraph(title, after=8)
    run = title.add_run("网络态势感知与可视化系统")
    set_run_font(run, size=24, bold=True, color=COLORS["navy"])
    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=14)
    run = subtitle.add_run("面向三层异构通信网络的拓扑建模、实时态势更新、交互分析与历史复盘原型")
    set_run_font(run, size=12, color=COLORS["muted"])

    rows = [
        ("资料用途", "供商业计划书、路演 PPT、答辩材料和技术可行性说明复用"),
        ("代码依据", "E:\\react-projects\\net-visualization 当前仓库源码"),
        ("整理日期", "2026-05-14"),
        ("当前形态", "可运行工程原型，含前端可视化、Node 后端、Python 数据导出与本地模拟数据"),
    ]
    add_table(doc, ["项目", "说明"], rows, [1700, 7660], header_fill=COLORS["blue_fill"], font_size=10)

    add_callout(
        doc,
        "一页结论",
        "项目已经形成“数据建模 - 后端接口 - 前端可视化 - 事件更新 - 历史回放”的原型闭环，适合作为创新大赛中的技术可行性证明和产品演示基础。后续重点应放在真实数据接入、WebSocket 通道完善、异常场景库、测试文档和商业化表达上。",
        fill="E8EEF5",
    )


def add_overview(doc):
    add_heading(doc, "1. 项目介绍", 1)
    add_para(
        doc,
        "本项目面向网络运维、规划与安全分析场景，构建一个三层异构通信网络态势感知与可视化原型。系统围绕骨干网、自组网、接入网等多层网络实体，对节点、链路、跨层关系、事件和态势快照进行统一建模，并在地图视图中呈现网络运行状态、链路质量、告警变化和历史回放。",
    )
    add_para(
        doc,
        "当前代码已经从“静态拓扑展示”推进到“动态态势原型”：前端提供 2D/3D 网络视图、节点列表、筛选搜索、详情弹窗、KPI 面板和回放控制；后端提供 REST API、事件存储、告警派生、连通性和最短路径分析；Python 模块负责生成三层网络样例、指标和事件 JSON，并通过命令队列接收前端节点增删改操作。",
    )
    add_heading(doc, "1.1 面向比赛的项目定位", 2)
    for item in [
        "目标用户：网络运维人员、网络规划人员、安全分析人员、系统管理员。",
        "核心价值：把多层异构网络的节点、链路、状态、告警和历史变化集中到一个可解释、可交互的态势视图中。",
        "演示重点：实时拓扑刷新、异常链路/节点可视化、跨层关系展示、路径追踪、历史复盘、Python 仿真数据驱动。",
        "比赛表达：可包装为“面向复杂通信网络的智能运维与可视化决策平台”原型。",
    ]:
        add_bullet(doc, item)


def add_architecture(doc):
    add_heading(doc, "2. 系统架构与数据流", 1)
    add_para(
        doc,
        "系统采用前后端分离和本地数据生成的轻量架构，符合学生项目快速演示、可验证和可扩展的目标。前端负责视觉编码和交互，Node 后端负责模型适配、API、事件和分析逻辑，Python 负责网络样例构建、指标计算和事件导出。",
    )
    rows = [
        ("React 前端", "src/App.js, src/Map3DView.js, src/NodeList.js", "地图渲染、筛选搜索、节点/链路详情、KPI、告警、历史回放、节点操作面板"),
        ("前端服务层", "src/services/networkApi.js, src/services/xduCampusPreset.js", "封装 REST 请求，映射校园场景节点，维护前端数据契约"),
        ("Node 后端", "server/src/server.js", "提供 REST API、统一响应 envelope、事件写入、告警派生、连通性/路径分析、回放帧缓存"),
        ("仓储适配层", "server/src/repositories/*.js", "支持内存种子数据和 Python JSON 导出数据两种来源"),
        ("Python 数据层", "Project-001/three_layer_network.py, live_export_runner.py", "构建三层网络、计算健康指标、输出 snapshot/metrics/events、处理命令队列"),
    ]
    add_table(doc, ["层次", "主要文件", "职责"], rows, [1500, 2800, 5060], header_fill=COLORS["blue_fill"], font_size=9.2)
    add_heading(doc, "2.1 当前数据流", 2)
    for step in [
        "Python 脚本生成 snapshot.json、metrics.json 和节点/链路事件 JSON。",
        "Node 后端读取 Python 导出文件，归一化字段并派生 crossLayerRelations。",
        "React 前端按默认 5 秒轮询 REST API，也保留本地 mock fallback。",
        "前端地图根据节点状态、链路质量、告警和回放帧进行可视化更新。",
        "前端节点增删改命令通过 POST /api/python/commands 写入 command_queue.jsonl，由 Python loop 消费。",
    ]:
        add_number(doc, step)


def add_features(doc):
    add_heading(doc, "3. 已实现功能", 1)
    rows = [
        ("统一数据模型", "已实现", "Node、Link、CrossLayerRelation 在前端模型、后端 seed、Python adapter 中均有对应结构；节点包含层级、位置、状态、指标；链路包含带宽、时延、丢包、可用性、SNR 等字段。"),
        ("REST API", "已实现", "支持 /api/topology、/api/nodes/:id、/api/links/:id、/api/situation/current、/api/events、/api/alerts、/api/playback/frames、/api/analysis/connectivity、/api/analysis/path、POST 事件与 Python 命令。"),
        ("实时/准实时更新", "部分实现", "前端主要通过轮询 REST 刷新，Python live loop 可每 3 秒重写导出数据；仓库中也有 WebSocket hook 和 mock server，但主流程尚未完全切到版本化 WebSocket channel。"),
        ("动态拓扑可视化", "已实现", "React Leaflet 2D 地图展示节点、链路、状态 halo、链路动画和 popup；支持地图飞行定位和链路状态颜色。"),
        ("3D 视图", "已实现", "Deck.gl + MapLibre 渲染带高度的节点、链路、文字标签和高度图例，可展示多层/高程差异。"),
        ("交互分析", "已实现", "支持搜索、层级过滤、节点列表分组、节点/链路点击详情、告警定位、连通性分析、按时延权重的最短路径分析。"),
        ("历史回放", "已实现", "后端维护 memory-ring 回放帧，前端提供进入回放、播放/暂停、上一帧/下一帧、时间轴滑块和 Live/Playback 状态条。"),
        ("事件与告警", "已实现", "支持用户事件写入、Python 事件映射、节点离线/链路质量阈值告警派生、Recent Alerts 和 Recent Events 面板。"),
        ("Python 集成", "已实现", "Python 构建三层网络样例、计算 networkHealth/onlineRate/avgDelay/avgLoss/avgLoad，支持 node add/remove/update 命令队列。"),
        ("校园场景映射", "已实现", "xduCampusPreset 将网络节点映射到校园楼宇、区域和 IoT 设备类型，更适合大赛演示的实景化表达。"),
    ]
    add_table(doc, ["功能模块", "状态", "说明"], rows, [1700, 1200, 6460], header_fill=COLORS["blue_fill"], font_size=8.8)


def add_stack(doc):
    add_heading(doc, "4. 技术栈", 1)
    rows = [
        ("前端框架", "React 19, React DOM, Create React App", "页面状态、组件组织、开发与构建"),
        ("地图与可视化", "Leaflet, React Leaflet, Deck.gl, MapLibre GL, React Map GL", "2D 地图、3D 场景、节点/链路线层、交互拾取"),
        ("UI 与图标", "Tailwind CSS, lucide-react, CSS Modules/普通 CSS", "控制面板、状态卡片、列表、图例、按钮图标"),
        ("后端运行时", "Node.js ESM, http, fs/path/crypto", "轻量 REST 服务、CORS、统一 JSON 响应、文件事件存储"),
        ("实时通信基础", "ws, WebSocket Hook", "已存在 mock WebSocket 服务和前端 hook，当前主链路以 REST polling 为主"),
        ("算法与仿真", "Python 3 dataclasses, heapq, JSON 文件导出", "三层网络建模、最短路径、网络健康指标、随机动态事件"),
        ("测试与工程", "Jest, React Testing Library, npm scripts", "具备基础测试依赖和脚本，现有测试用例仍停留在 CRA 默认模板"),
    ]
    add_table(doc, ["类别", "技术/依赖", "用途"], rows, [1600, 3000, 4760], header_fill=COLORS["blue_fill"], font_size=9)


def add_models_and_api(doc):
    add_heading(doc, "5. 数据模型与接口摘要", 1)
    add_heading(doc, "5.1 核心实体", 2)
    rows = [
        ("Node", "id、name、type、layer、location.geo、state、metrics、role、alarmLevel", "表示路由器、基站、自组网节点、终端、卫星或校园 IoT 设备。"),
        ("Link", "id、from、to、type、bandwidthMbps、delayMs、lossRate、snrDb、utilization、availability、health/state", "表示有线或无线链路质量，供颜色编码、告警和路径分析使用。"),
        ("CrossLayerRelation", "id、fromNodeId、toNodeId、relationType、notes", "表示 access/backhaul/relay 等跨层依赖关系，后端可根据链路自动派生。"),
        ("NetworkEvent", "id、type、severity、entityType、entityId、message、status、occurredAt、payload", "表示节点/链路状态变化、手工事件和 Python 导出事件。"),
        ("SituationSnapshot", "snapshotAt、nodeSummary、linkSummary、alarmSummary、healthScore、pythonMetrics", "表示当前态势汇总，用于 KPI 和健康评分。"),
        ("MetricSeries", "前端 kpiHistory 数组", "保存在线节点、告警、平均时延、丢包率、利用率等趋势点。"),
    ]
    add_table(doc, ["实体", "主要字段", "说明"], rows, [1700, 3500, 4160], header_fill=COLORS["blue_fill"], font_size=8.7)

    add_heading(doc, "5.2 API 清单", 2)
    rows = [
        ("GET", "/api/topology", "获取当前拓扑和跨层关系"),
        ("GET", "/api/nodes/:id", "获取单个节点详情"),
        ("GET", "/api/links/:id", "获取单条链路详情"),
        ("GET", "/api/situation/current", "获取当前态势评分和汇总指标"),
        ("GET", "/api/events?limit=", "获取事件列表"),
        ("GET", "/api/alerts?active=true", "获取由节点/链路状态派生的告警"),
        ("GET", "/api/playback/frames?limit=", "获取历史回放帧"),
        ("GET", "/api/analysis/connectivity", "获取连通性、连通分量和孤立节点"),
        ("GET", "/api/analysis/path?from=&to=", "获取按时延权重的最短路径"),
        ("POST", "/api/topology/events", "写入手工拓扑事件"),
        ("POST", "/api/python/commands", "写入 Python 命令队列，支持 node:add、node:remove、node:update"),
    ]
    add_table(doc, ["方法", "路径", "作用"], rows, [900, 3100, 5360], header_fill=COLORS["blue_fill"], font_size=8.8)


def add_shortcomings(doc):
    add_heading(doc, "6. 当前不足", 1)
    rows = [
        ("实时通道尚未完全闭环", "WebSocket hook 和 mock server 已存在，但主应用主要依赖 REST polling；尚未完成 topology:update、metric:update、alarm:update、snapshot:update 的正式版本化通道。"),
        ("持久化能力偏原型", "事件使用 JSON 文件或内存缓存，回放帧是 memory-ring；尚未接入数据库、索引、长期历史查询和数据清理策略。"),
        ("测试覆盖不足", "仓库仍保留 CRA 默认 App.test.js，未覆盖地图渲染、REST API、Python adapter、回放和分析算法。"),
        ("真实网络适配不足", "当前主要使用模拟/导出数据，尚未接入真实 SNMP、NetFlow、日志、设备遥测或安全告警源。"),
        ("性能验证不足", "没有针对中等规模拓扑的压力测试和 FPS/接口耗时报告，30 FPS 和历史查询目标尚缺实测数据。"),
        ("安全与部署较弱", "没有生产级认证、权限、审计、容器化部署、配置管理和异常恢复策略。按项目范围这是合理取舍，但参赛材料需说明未来计划。"),
        ("文档和编码需清理", "部分中文 README/注释出现编码乱码，接口契约、数据字典、部署说明还需要整理成正式工程文档。"),
        ("智能分析仍是扩展项", "当前已具备连通性和最短路径，尚未实现根因定位、预测、数字孪生或自动优化。"),
    ]
    add_table(doc, ["不足项", "说明"], rows, [2200, 7160], header_fill="FFF2CC", font_size=9)


def add_roadmap(doc):
    add_heading(doc, "7. 后续计划", 1)
    rows = [
        ("第 1 阶段：比赛演示增强", "1-2 周", "修复文档乱码；补齐启动说明；制作稳定演示脚本；完善真实演示数据集；为核心场景增加一键导入和回放样例。"),
        ("第 2 阶段：实时链路闭环", "2-4 周", "将 WebSocket 正式接入主应用，定义 versioned envelope，支持 topology:update、metric:update、alarm:update、snapshot:update；断线重连后按 last-known-point 补偿。"),
        ("第 3 阶段：工程化验证", "3-5 周", "补充 API 单元测试、React 组件测试、Python adapter 测试、E2E 演示测试；输出接口契约和验收清单。"),
        ("第 4 阶段：数据与算法升级", "4-8 周", "引入数据库或时序存储；沉淀异常场景库；扩展健康评分、趋势分析、告警关联和跨层影响分析。"),
        ("第 5 阶段：产品化包装", "8 周以后", "优化部署、权限、安全审计和多用户协作；对接真实设备或仿真平台；探索根因分析、预测维护和数字孪生能力。"),
    ]
    add_table(doc, ["阶段", "时间建议", "重点任务"], rows, [2200, 1200, 5960], header_fill=COLORS["blue_fill"], font_size=8.8)

    add_heading(doc, "7.1 队友写商业计划书/PPT 可复用的表达", 2)
    for item in [
        "痛点：多层异构网络设备类型多、状态分散、故障定位依赖人工经验，缺少统一态势视图和历史复盘能力。",
        "方案：构建统一模型和可视化平台，把节点、链路、告警、指标、跨层关系和历史变化汇聚到一张可交互态势图中。",
        "差异点：兼顾 2D 地理视图、3D 高程视图、跨层追踪、事件驱动更新和 Python 仿真数据闭环，便于演示和教学科研扩展。",
        "落地场景：校园/园区 IoT 网络、应急通信、无人机自组网、边缘接入网络、专网运维保障。",
        "阶段成果：已有可运行原型，可支撑现场演示、评委答辩和后续工程迭代。",
    ]:
        add_bullet(doc, item)


def add_appendix(doc):
    add_heading(doc, "8. 运行与验证建议", 1)
    add_heading(doc, "8.1 推荐启动顺序", 2)
    for step in [
        "执行 npm install 安装前端依赖，并在 server 目录安装后端依赖。",
        "执行 npm run generate:python 生成 Python 初始数据。",
        "执行 npm run start:server 启动 Node 后端，默认地址 http://localhost:8080。",
        "执行 npm start 启动 React 前端，默认地址 http://localhost:3000。",
        "如需准实时数据，另开终端执行 npm run generate:python:loop。",
    ]:
        add_number(doc, step)

    add_heading(doc, "8.2 验收检查清单", 2)
    for item in [
        "地图可以看到多层节点和链路，节点点击后弹出详情。",
        "筛选、搜索、节点列表分组、图例折叠和地图定位可正常使用。",
        "右侧控制面板显示当前健康分、在线节点、告警和事件。",
        "进入 Playback 后可以播放、暂停、拖动历史帧，并能回到 Live 模式。",
        "连通性分析和最短路径分析能返回结果并定位路径节点。",
        "Python loop 运行时，节点/链路状态和指标能随时间刷新。",
        "新增/删除节点命令能写入 command_queue.jsonl，并被 Python loop 消费。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.3 关键源码索引", 2)
    rows = [
        ("前端主应用", "src/App.js"),
        ("3D 视图", "src/Map3DView.js"),
        ("节点列表", "src/NodeList.js"),
        ("REST API 客户端", "src/services/networkApi.js"),
        ("校园场景映射", "src/services/xduCampusPreset.js"),
        ("后端 API", "server/src/server.js"),
        ("Python 适配仓储", "server/src/repositories/pythonFileNetworkRepository.js"),
        ("内存仓储", "server/src/repositories/inMemoryNetworkRepository.js"),
        ("Python 三层网络模型", "Project-001/three_layer_network.py"),
        ("Python 实时导出循环", "Project-001/live_export_runner.py"),
    ]
    add_table(doc, ["用途", "文件"], rows, [2600, 6760], header_fill=COLORS["blue_fill"], font_size=9)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_overview(doc)
    add_architecture(doc)
    add_features(doc)
    add_stack(doc)
    add_models_and_api(doc)
    add_shortcomings(doc)
    add_roadmap(doc)
    add_appendix(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
