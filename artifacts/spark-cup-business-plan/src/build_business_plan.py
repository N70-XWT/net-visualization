from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
WORK = ROOT / "artifacts" / "spark-cup-business-plan"
ASSETS = WORK / "assets"
OUT = WORK / "output"
RENDER = WORK / "rendered"
DECK_ASSETS = ROOT / "artifacts" / "spark-cup-deck" / "assets"

ASSETS.mkdir(parents=True, exist_ok=True)
OUT.mkdir(parents=True, exist_ok=True)
RENDER.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT / "物联慧眼-星火杯版本-商业计划书.docx"
HERO = ASSETS / "network-ops-hero.png"
SCREENSHOT = ASSETS / "product-screenshot.png"
FINANCE_CHART = ASSETS / "finance_chart.png"

if (DECK_ASSETS / "network-ops-hero.png").exists():
    shutil.copyfile(DECK_ASSETS / "network-ops-hero.png", HERO)
if (DECK_ASSETS / "product-screenshot.png").exists():
    shutil.copyfile(DECK_ASSETS / "product-screenshot.png", SCREENSHOT)

with open(ROOT / "Project-001" / "metrics.json", "r", encoding="utf-8") as f:
    metrics = json.load(f)


NAVY = "0B1524"
NAVY2 = "12233A"
BLUE = "1E88E5"
CYAN = "19C7E8"
GREEN = "16C784"
AMBER = "F5B84B"
RED = "E95B67"
LIGHT = "F6F9FC"
MUTED = "5C6F82"
BORDER = "DCE6F0"
WHITE = "FFFFFF"


def font_path(name: str) -> str:
    path = Path("C:/Windows/Fonts") / name
    return str(path) if path.exists() else "arial.ttf"


def make_finance_chart(path: Path) -> None:
    width, height = 1400, 650
    img = Image.new("RGB", (width, height), "#F6F9FC")
    draw = ImageDraw.Draw(img)
    font_regular = ImageFont.truetype(font_path("msyh.ttc"), 30)
    font_small = ImageFont.truetype(font_path("msyh.ttc"), 22)
    font_bold = ImageFont.truetype(font_path("msyhbd.ttc"), 34)
    title_font = ImageFont.truetype(font_path("msyhbd.ttc"), 42)

    draw.text((56, 40), "三年收入测算（万元）", fill="#0B1524", font=title_font)
    draw.text((58, 100), "测算口径：PoC 试点 + 私有化部署 + 模块授权与年度服务", fill="#5C6F82", font=font_regular)

    left, top, bottom = 110, 190, 540
    chart_w = 1000
    max_val = 240
    years = ["2026", "2027", "2028"]
    project = [18, 56, 96]
    service = [6, 42, 128]
    colors = ["#16C784", "#19C7E8"]

    for i in range(0, 7):
        y = bottom - i * (bottom - top) / 6
        draw.line((left, y, left + chart_w, y), fill="#DCE6F0", width=2)
        draw.text((45, y - 13), str(i * 40), fill="#5C6F82", font=font_small)

    group_gap = chart_w / 3
    bar_w = 88
    for i, year in enumerate(years):
        cx = left + group_gap * i + group_gap / 2
        vals = [project[i], service[i]]
        for j, val in enumerate(vals):
            h = (val / max_val) * (bottom - top)
            x0 = cx - bar_w - 16 + j * (bar_w + 32)
            y0 = bottom - h
            draw.rounded_rectangle((x0, y0, x0 + bar_w, bottom), radius=18, fill=colors[j])
            draw.text((x0 + 20, y0 - 34), str(val), fill="#0B1524", font=font_small)
        draw.text((cx - 42, bottom + 28), year, fill="#0B1524", font=font_regular)

    draw.rounded_rectangle((1150, 238, 1310, 282), radius=10, fill=colors[0])
    draw.text((1330, 238), "项目制收入", fill="#0B1524", font=font_small)
    draw.rounded_rectangle((1150, 310, 1310, 354), radius=10, fill=colors[1])
    draw.text((1330, 310), "授权与服务", fill="#0B1524", font=font_small)
    draw.text((1150, 430), "2028 年收入测算", fill="#5C6F82", font=font_small)
    draw.text((1150, 468), "224 万元", fill="#0B1524", font=font_bold)
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold=False, color="1F2937", size=9.5, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_table(table, header_fill=NAVY2, header_text=WHITE, body_fill="FFFFFF") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_shading(cell, header_fill if row_idx == 0 else body_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
            if row_idx == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor.from_string(header_text)
                        r.bold = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text, bold=False, color="1F2937", size=10.5):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_paragraph(doc, text="", style=None, color="1F2937", size=10.5, bold=False, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    if text:
        add_run(p, text, bold=bold, color=color, size=size)
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Heading {}".format(level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p.add_run(text)


def add_note_box(doc, title, body, fill="EAF8FF", accent=CYAN):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [0.28, 15.7])
    cell0, cell1 = table.rows[0].cells
    set_cell_shading(cell0, accent)
    set_cell_shading(cell1, fill)
    set_cell_border(cell0, accent)
    set_cell_border(cell1, fill)
    p = cell1.paragraphs[0]
    add_run(p, title, bold=True, color=NAVY, size=10.5)
    p.add_run("\n")
    add_run(p, body, color="263445", size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(doc, text, level=0, color="1F2937"):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    add_run(p, text, color=color, size=10.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18


def add_caption(doc, text):
    p = add_paragraph(doc, text, color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return p


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.85)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")

    for name, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, NAVY),
        ("Heading 3", 12, BLUE),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "物联慧眼商业计划书 | 星火杯版本 | 参赛信息待完善", color=MUTED, size=8.5)


def add_cover(doc: Document) -> None:
    if HERO.exists():
        doc.add_picture(str(HERO), width=Cm(16.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_paragraph(doc, "物联慧眼", size=32, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=2)
    p = add_paragraph(doc, "三层异构网络态势感知与可视化平台", size=17, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_paragraph(doc, "商业计划书（星火杯版本）", size=14, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    info = doc.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    set_table_width(info, [4.2, 10.8])
    fields = [
        ("项目名称", "物联慧眼：三层异构网络态势感知与可视化平台"),
        ("提交版本", "基于 Git 提交 ee715bc《星火杯物联慧眼》"),
        ("参赛赛道/组别", "待填写（建议结合校赛通知选择高教主赛道创意组或创业组）"),
        ("负责人/团队", "待填写"),
        ("联系方式", "手机/邮箱/QQ 待填写"),
    ]
    for idx, (k, v) in enumerate(fields):
        set_cell_text(info.rows[idx].cells[0], k, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(info.rows[idx].cells[1], v, color="1F2937", size=9.5)
        set_cell_shading(info.rows[idx].cells[0], NAVY2)
        set_cell_shading(info.rows[idx].cells[1], "F8FBFF")
        set_cell_border(info.rows[idx].cells[0])
        set_cell_border(info.rows[idx].cells[1])
    add_paragraph(doc, "日期：2026 年 5 月 | 本稿为参赛包装版，市场调研、团队信息、知识产权与试点证明需赛前补充。", color=MUTED, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, before=12)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    toc_items = [
        "1. 执行摘要",
        "2. 赛事评审对齐与项目定位",
        "3. 行业痛点与市场机会",
        "4. 产品与解决方案",
        "5. 技术路线与研发基础",
        "6. 核心创新与竞争壁垒",
        "7. 目标客户、商业模式与定价",
        "8. 市场推广与运营计划",
        "9. 财务测算与融资规划",
        "10. 团队建设与资源配置",
        "11. 风险识别与应对",
        "12. 附录：当前版本证据与参考资料",
    ]
    for item in toc_items:
        add_paragraph(doc, item, color=NAVY, size=11.2, after=5)
    add_note_box(
        doc,
        "阅读提示",
        "本商业计划书定位为创新创业大赛参赛材料：正文突出项目商业价值、技术创新、落地路径与团队成长；其中真实客户、市场规模、财务收入、团队资质等内容需以赛前调研材料最终校准。",
        fill="FFF8E6",
        accent=AMBER,
    )
    doc.add_page_break()


def add_exec_summary(doc: Document) -> None:
    add_heading(doc, "1. 执行摘要", 1)
    add_paragraph(
        doc,
        "物联慧眼是一套面向三层异构通信网络的态势感知与可视化平台，聚焦骨干网络、Ad hoc 自组织网络、接入网络及终端设备之间的统一建模、实时更新、交互分析和历史复盘。项目当前已形成可运行原型，具备 REST API、动态拓扑展示、告警推导、历史帧回放、连通性分析和时延加权最短路径分析等能力。",
    )
    add_note_box(
        doc,
        "一句话定位",
        "将割裂的网络设备、链路质量、跨层关系和事件快照整合为一个可视、可算、可复盘的网络态势操作平台。",
        fill="EAF8FF",
        accent=CYAN,
    )
    table = doc.add_table(rows=5, cols=4)
    table.autofit = False
    set_table_width(table, [3.1, 4.5, 4.5, 4.0])
    headers = ["维度", "当前基础", "商业包装表达", "赛事价值"]
    rows = [
        ["产品", "React + Leaflet 前端、Node.js 后端、Python 场景数据", "三层异构网络态势感知平台", "可演示、可部署、可扩展"],
        ["技术", "统一 Node/Link/Relation/Event/Snapshot 模型", "跨层建模 + 事件驱动 + 历史回放", "项目创新与专业能力"],
        ["客户", "校园/园区、应急通信、工业物联、边缘网络", "先试点再行业复制", "产业价值与落地路径"],
        ["收入", "PoC、私有化部署、模块授权、年度服务", "轻资产项目制起步，逐步标准化", "商业模式清晰"],
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, color="1F2937", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if c == 0 else None)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_paragraph(doc, "当前演示指标：节点 19 个、链路 23 条、网络健康度 69%、平均时延 55.8ms、平均丢包率 4.29%。这些指标来自 Project-001/metrics.json，可作为原型可运行性的量化证据。", color=MUTED, size=9.5)


def add_competition_fit(doc: Document) -> None:
    add_heading(doc, "2. 赛事评审对齐与项目定位", 1)
    add_paragraph(
        doc,
        "中国国际大学生创新创业大赛/中国国际大学生创新大赛强调学生成长、项目创新、产业价值和团队协作。本项目适合按照“新工科 + 信息通信 + 网络安全/物联网运维”的方向包装，以技术原型支撑商业计划，以真实行业痛点支撑市场逻辑。",
    )
    table = doc.add_table(rows=5, cols=3)
    table.autofit = False
    set_table_width(table, [3.2, 6.0, 6.7])
    data = [
        ["评审关注", "材料表达重点", "本项目对应内容"],
        ["个人成长", "专业知识与商业知识结合、调研深入、团队能力提升", "通信网络、Web 工程、数据建模、可视化交互、商业测算"],
        ["项目创新", "问题导向、目标导向、创新成效", "跨层关系一等建模、事件快照、历史回放、可计算拓扑分析"],
        ["产业价值", "市场定位、商业模式、发展前景", "面向应急通信、校园/园区专网、工业物联与边缘网络运维"],
        ["团队协作", "组织结构、资源整合、创业可能性", "待补充成员分工、导师资源、实验平台、试点单位与知识产权"],
    ]
    for r, row in enumerate(data):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.8)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_note_box(
        doc,
        "参赛包装建议",
        "若团队尚未注册企业或没有销售订单，建议以“高教主赛道创意组/成长组”语言包装：强调原型可运行、应用场景明确、成果转化路径可信；若已有试点合同或企业主体，可进一步按创业组口径强化经营数据。",
        fill="F2FFF8",
        accent=GREEN,
    )


def add_problem_market(doc: Document) -> None:
    add_heading(doc, "3. 行业痛点与市场机会", 1)
    add_paragraph(
        doc,
        "专网、应急通信、校园/园区物联网和边缘网络正在从单一链路、单一设备监控，转向跨层、多源、动态协同的网络态势管理。传统网管工具重资产、部署门槛高，GIS 或地图系统更偏位置呈现，IoT 平台偏设备接入和业务数据，对“网络拓扑关系、链路质量、告警事件、历史回放”的联动支持不足。",
    )
    for point in [
        "态势割裂：骨干、接入、无线中继、终端与卫星/无人机回传分属不同系统，故障定位需要人工拼图。",
        "事件滞后：链路丢包、时延、SNR、节点离线等变化不能及时进入统一事件流，影响应急响应。",
        "复盘缺口：演练和事故后缺少可回放的历史快照，难以复现故障传播路径和处置过程。",
        "成本矛盾：中小型园区、实验室和教学科研场景需要轻量可部署工具，而非重型商业网管。",
    ]:
        add_bullet(doc, point)
    table = doc.add_table(rows=5, cols=4)
    table.autofit = False
    set_table_width(table, [3.0, 4.8, 4.5, 3.7])
    rows = [
        ["场景", "典型需求", "付费动因", "切入优先级"],
        ["高校/园区专网", "IoT、安防、边缘节点、实验网络统一展示", "低成本演示、运维效率、科研平台", "高"],
        ["应急通信演练", "临时组网、无人机中继、卫星回传态势", "快速定位、演练复盘、指挥可视化", "高"],
        ["工业物联网", "远程站点链路质量与资产状态监控", "停机损失控制、运维降本", "中"],
        ["运营商边缘网络", "多接入边缘链路与终端拓扑联动", "系统集成、模块化补强", "中"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER if c in [0, 3] else None)
    style_table(table)
    set_repeat_table_header(table.rows[0])


def add_product_solution(doc: Document) -> None:
    add_heading(doc, "4. 产品与解决方案", 1)
    add_paragraph(
        doc,
        "物联慧眼以“统一建模 - 实时更新 - 可视分析 - 历史复盘”为核心闭环。前端负责地图渲染、节点筛选、链路状态表达、详情面板、回放控制和分析结果展示；后端负责模型持久化、Python 场景数据适配、事件处理、告警推导、REST API 与历史帧组织。",
    )
    if SCREENSHOT.exists():
        doc.add_picture(str(SCREENSHOT), width=Cm(15.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "图 1 当前星火杯版本产品界面截图：地图态势、控制面板、KPI、回放与告警入口")
    table = doc.add_table(rows=6, cols=3)
    table.autofit = False
    set_table_width(table, [3.2, 6.3, 6.5])
    rows = [
        ["产品模块", "核心功能", "客户价值"],
        ["统一拓扑建模", "Node、Link、CrossLayerRelation、NetworkEvent、SituationSnapshot", "将跨层通信关系变成可计算资产"],
        ["实时态势看板", "动态节点、链路健康、KPI、告警、过滤搜索、详情联动", "减少人工切换系统和定位时间"],
        ["历史回放", "内存环形帧缓存、事件时间轴、播放/暂停/帧切换", "支持应急演练复盘和责任分析"],
        ["图分析能力", "连通性分析、时延加权最短路径、影响范围提示", "从“看到异常”走向“理解影响”"],
        ["场景适配", "Python 快照/指标/事件文件接入，命令队列支持动态节点操作", "便于教学科研、行业演练和二次开发"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.7)
    style_table(table)
    set_repeat_table_header(table.rows[0])


def add_tech(doc: Document) -> None:
    add_heading(doc, "5. 技术路线与研发基础", 1)
    add_paragraph(
        doc,
        "项目采用 React + Leaflet/MapLibre/Deck.gl 的前端可视化技术栈，后端采用 Node.js 原生 HTTP 服务与仓储模式，Python 侧提供三层网络场景数据生成与导出。该架构符合原型阶段“可运行、可演示、可扩展”的要求，同时保留未来对真实采集系统、数据库和 WebSocket 通道的集成空间。",
    )
    table = doc.add_table(rows=5, cols=4)
    table.autofit = False
    set_table_width(table, [3.0, 4.0, 4.7, 4.3])
    rows = [
        ["层级", "实现基础", "当前证据", "扩展方向"],
        ["前端可视层", "React、Leaflet、MapLibre、Deck.gl", "2D/3D 模式、节点列表、过滤搜索、详情、回放面板", "大屏模式、告警联动、趋势图"],
        ["服务接口层", "Node.js REST API", "/api/topology、/api/events、/api/playback/frames 等", "WebSocket 频道、API 鉴权、接口文档"],
        ["数据模型层", "统一拓扑仓储与 Python 文件适配", "Node/Link/Relation/Event/Snapshot 统一组织", "数据库持久化、增量同步"],
        ["分析计算层", "连通性、最短路径、告警阈值", "delay/loss/utilization/SNR 告警推导", "根因分析、趋势预测、仿真评估"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.5)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_note_box(
        doc,
        "工程可信度",
        "星火杯版本已通过 npm run build 构建验证；当前演示数据包含 19 个节点、23 条链路，能够展示离线节点、链路劣化、历史帧与实时 KPI。",
        fill="EAF8FF",
        accent=CYAN,
    )


def add_innovation(doc: Document) -> None:
    add_heading(doc, "6. 核心创新与竞争壁垒", 1)
    innovations = [
        ("跨层关系一等建模", "CrossLayerRelation 不再只是 UI 连线，而是进入数据模型、接口返回和分析逻辑，可支持从终端到中继、接入、骨干的链路追踪。"),
        ("事件驱动态势快照", "拓扑、指标、告警与事件通过时间戳组织，为历史回放、演练复盘和事故审计提供统一依据。"),
        ("可计算网络态势", "连通性分析、时延加权最短路径和告警阈值推导，使系统具备“看见 - 判断 - 分析”的闭环能力。"),
        ("轻量化场景仿真", "Python 数据生成器与命令队列支持节点新增、移除、状态更新等场景，便于校赛答辩和客户试点演示。"),
    ]
    table = doc.add_table(rows=1 + len(innovations), cols=3)
    table.autofit = False
    set_table_width(table, [3.6, 8.0, 4.0])
    headers = ["创新点", "说明", "壁垒来源"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, (title, desc) in enumerate(innovations, start=1):
        set_cell_text(table.rows[r].cells[0], title, bold=True, color=NAVY, size=8.8)
        set_cell_text(table.rows[r].cells[1], desc, color="1F2937", size=8.6)
        set_cell_text(table.rows[r].cells[2], "数据模型 + 工程实现 + 场景经验", color=GREEN, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_paragraph(
        doc,
        "知识产权建议：赛前可将“多层异构网络态势感知系统”“历史回放与跨层关系分析方法”“网络拓扑动态演示平台”等方向整理为软著材料，并保留代码提交记录、演示视频、设计文档作为成果证明。",
        color=MUTED,
        size=9.5,
    )


def add_business_model(doc: Document) -> None:
    add_heading(doc, "7. 目标客户、商业模式与定价", 1)
    add_paragraph(
        doc,
        "项目商业化不宜一开始定位为重型商业网管，而应从“轻量可部署的态势演示/运维分析平台”切入。以高校、实验室、园区和应急演练项目获取种子案例，再逐步沉淀行业模板和可复用模块。",
    )
    table = doc.add_table(rows=4, cols=4)
    table.autofit = False
    set_table_width(table, [3.0, 4.4, 4.4, 4.2])
    rows = [
        ["收入线", "收费方式", "典型报价建议", "交付内容"],
        ["PoC 试点", "项目制", "3-8 万元/项目", "场景搭建、演示大屏、基础培训"],
        ["私有化部署", "软件 + 实施", "10-30 万元/套", "系统部署、数据适配、接口联调"],
        ["年度服务", "订阅/维保", "合同额 15%-25%/年", "运维支持、版本更新、演练脚本"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER if c in [0, 2] else None)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_note_box(
        doc,
        "商业模式画布",
        "关键伙伴：高校实验室、系统集成商、应急演练单位；关键活动：数据适配、场景封装、可视分析算法；价值主张：复杂网络看得见、找得快、可复盘；客户关系：试点共创 + 年度服务；渠道：校企合作、创新创业赛事、实验室示范、集成商项目。",
        fill="F2FFF8",
        accent=GREEN,
    )


def add_competition(doc: Document) -> None:
    add_heading(doc, "8. 市场推广与运营计划", 1)
    add_paragraph(
        doc,
        "推广策略采用“比赛曝光 - 校园试点 - 行业演练 - 集成合作”的路径。比赛阶段强调原型演示和商业计划完整性；试点阶段重点获取真实反馈；商业阶段以场景包和接口适配降低交付成本。",
    )
    phases = [
        ["阶段", "时间", "重点目标", "关键成果"],
        ["赛事版", "2026 Q2-Q3", "完善计划书、路演稿、演示视频和答辩材料", "进入校赛/省赛、形成原型证明"],
        ["试点版", "2026 Q4-2027 Q2", "完成 1-3 个校园或实验室场景试点", "试用证明、数据样例、软著"],
        ["行业版", "2027 Q3-2028", "面向应急通信/园区专网复制演练方案", "行业场景包、标准报价、合作伙伴"],
    ]
    table = doc.add_table(rows=len(phases), cols=4)
    table.autofit = False
    set_table_width(table, [2.5, 3.1, 5.4, 5.0])
    for r, row in enumerate(phases):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.6, align=WD_ALIGN_PARAGRAPH.CENTER if c in [0, 1] else None)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_heading(doc, "8.1 竞品与差异化", 2)
    rows = [
        ["对比对象", "优势", "不足", "物联慧眼切入点"],
        ["传统网管系统", "监控能力强、成熟稳定", "成本高、重部署、定制慢", "轻量原型、跨层回放、演练导向"],
        ["GIS/地图平台", "空间展示能力强", "网络模型和事件逻辑弱", "把地理展示与网络拓扑分析结合"],
        ["IoT 平台", "设备接入和数据采集成熟", "网络链路和跨层关系表达不足", "补齐通信态势与运维复盘能力"],
        ["开源可视化组件", "灵活、成本低", "缺少完整业务闭环", "形成模型、接口、分析、回放完整方案"],
    ]
    table2 = doc.add_table(rows=len(rows), cols=4)
    table2.autofit = False
    set_table_width(table2, [2.8, 4.0, 4.0, 5.2])
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table2.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.4)
    style_table(table2)
    set_repeat_table_header(table2.rows[0])


def add_finance(doc: Document) -> None:
    add_heading(doc, "9. 财务测算与融资规划", 1)
    make_finance_chart(FINANCE_CHART)
    doc.add_picture(str(FINANCE_CHART), width=Cm(15.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 2 三年收入测算示意（单位：万元，赛前需结合真实报价和试点资源校准）")
    rows = [
        ["项目", "2026", "2027", "2028", "测算依据"],
        ["项目制收入", "18", "56", "96", "PoC 试点与私有化部署逐步增加"],
        ["授权与服务收入", "6", "42", "128", "模块授权、年度服务、场景包复制"],
        ["营业收入合计", "24", "98", "224", "以轻资产软件服务为主"],
        ["研发与交付成本", "18", "45", "70", "人员、测试、部署、适配"],
        ["销售与运营成本", "4", "18", "40", "宣传、差旅、试点支持"],
        ["云资源与软硬件", "3", "12", "24", "服务器、演示设备、测试环境"],
        ["预计净利润", "-1", "23", "90", "标准化后毛利率提升"],
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.autofit = False
    set_table_width(table, [3.1, 2.0, 2.0, 2.0, 6.8])
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0 or r == 3), color=WHITE if r == 0 else "1F2937", size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER if c in [1, 2, 3] else None)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_note_box(
        doc,
        "融资/资源需求",
        "若进入省赛或后续孵化，可提出 20-50 万元种子支持需求，主要用于真实数据接入、行业调研、软著/专利申请、试点部署和演示设备。资金需求需结合学校政策和团队计划最终调整。",
        fill="FFF8E6",
        accent=AMBER,
    )


def add_team_risk_appendix(doc: Document) -> None:
    add_heading(doc, "10. 团队建设与资源配置", 1)
    rows = [
        ["角色", "建议成员", "职责", "赛前补充材料"],
        ["项目负责人", "待填写", "商业统筹、答辩、外部沟通", "个人简介、获奖、项目经历"],
        ["前端与交互", "待填写", "地图可视化、控制面板、用户体验", "代码贡献、截图、演示视频"],
        ["后端与数据", "待填写", "REST API、仓储模型、事件和回放", "接口文档、测试记录"],
        ["算法与仿真", "待填写", "Python 场景生成、链路质量、分析算法", "模型说明、实验数据"],
        ["商业与运营", "待填写", "市场调研、竞品分析、财务测算", "访谈记录、报价表、推广计划"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.autofit = False
    set_table_width(table, [3.0, 3.0, 5.2, 4.7])
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.rows[r].cells[c], text, bold=(r == 0 or c == 0), color=WHITE if r == 0 else "1F2937", size=8.4)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    add_heading(doc, "11. 风险识别与应对", 1)
    risks = [
        ["市场验证不足", "缺少真实客户访谈和试点证明", "赛前完成 5-8 份访谈记录，争取实验室/学院试用意见"],
        ["数据接入复杂", "不同设备协议和字段差异大", "先做适配层和标准数据模型，逐步接入真实采集器"],
        ["商业化周期长", "政企/园区客户采购流程较长", "以教学科研和演练项目切入，降低首单门槛"],
        ["团队资源有限", "学生团队研发、销售、交付能力有限", "明确分工，引入导师、集成商或实验室资源"],
        ["竞品替代风险", "大型网管厂商能力更成熟", "聚焦轻量跨层态势、演练回放和二次开发差异化"],
    ]
    table2 = doc.add_table(rows=1 + len(risks), cols=3)
    table2.autofit = False
    set_table_width(table2, [3.6, 5.6, 6.7])
    for i, h in enumerate(["风险", "表现", "应对措施"]):
        set_cell_text(table2.rows[0].cells[i], h, bold=True, color=WHITE, size=9)
    for r, row in enumerate(risks, start=1):
        for c, text in enumerate(row):
            set_cell_text(table2.rows[r].cells[c], text, bold=(c == 0), color="1F2937", size=8.4)
    style_table(table2)
    set_repeat_table_header(table2.rows[0])

    add_heading(doc, "12. 附录：当前版本证据与参考资料", 1)
    add_heading(doc, "12.1 当前版本功能证据", 2)
    for item in [
        "Git 提交：ee715bc《星火杯物联慧眼》。",
        "前端构建：npm run build 已通过，产物位于 build/。",
        "核心接口：/api/topology、/api/nodes/:id、/api/links/:id、/api/situation/current、/api/events、/api/playback/frames、/api/analysis/connectivity、/api/analysis/path。",
        "当前数据：节点 19 个、链路 23 条、网络健康度 69%、在线率 73.68%、平均时延 55.8022ms、平均丢包率 4.29%。",
        "关键能力：动态拓扑、过滤搜索、节点详情、告警推导、历史帧回放、最短路径、Python 场景数据接入。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.2 参考资料", 2)
    refs = [
        "中国国际大学生创新大赛（2025）评审规则：公开材料显示高教主赛道关注个人成长、项目创新、产业价值、团队协作等维度。",
        "华中科技大学关于中国国际大学生创新大赛（2026）项目预征集的通知：公开材料列示项目计划书和路演 PPT 常见提交项，具体要求以所在学校当年通知为准。",
        "项目本地材料：AGENTS.md、server/src/server.js、src/App.js、Project-001/snapshot.json、Project-001/metrics.json。",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    add_note_box(
        doc,
        "最终提交前检查清单",
        "1. 补齐参赛人和团队真实信息；2. 增加市场访谈与试点证明；3. 将财务测算替换为团队确认口径；4. 附软著/论文/获奖/指导教师资源；5. 对照当年学校通知调整页数、格式和附件要求。",
        fill="FFF8E6",
        accent=AMBER,
    )


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_exec_summary(doc)
    add_competition_fit(doc)
    add_problem_market(doc)
    add_product_solution(doc)
    add_tech(doc)
    add_innovation(doc)
    add_business_model(doc)
    add_competition(doc)
    add_finance(doc)
    add_team_risk_appendix(doc)
    doc.core_properties.title = "物联慧眼：星火杯版本商业计划书"
    doc.core_properties.subject = "中国国际大学生创新创业大赛商业计划书"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
